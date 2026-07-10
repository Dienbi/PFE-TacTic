from docx import Document
import logging
from typing import Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class DocxTextExtractor:
    """Service for extracting text from DOCX files."""

    async def extract(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"

            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise

    def validate_file(self, file_path: str) -> bool:
        """Validate that the file is a valid DOCX."""
        if not Path(file_path).exists():
            return False

        try:
            Document(file_path)
            return True
        except Exception:
            return False
